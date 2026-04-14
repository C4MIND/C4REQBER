import uuid
import hashlib
from pathlib import Path
from typing import List, Optional, Dict
from dataclasses import dataclass


@dataclass
class Document:
    id: str
    title: str
    source_path: str
    chunks: int
    metadata: Dict


class DocumentIngester:
    def __init__(self):
        from .embedder import encode

        self.encode = encode
        self.storage_root = Path.home() / ".turbo-cdi" / "documents"
        self.storage_root.mkdir(parents=True, exist_ok=True)

    def ingest(self, file_path: Path, user_id: str = "default") -> Document:
        # Check for deduplication
        content_bytes = file_path.read_bytes()
        content_hash = hashlib.sha256(content_bytes).hexdigest()[:16]

        from .vector_store import UserDocumentStore

        store = UserDocumentStore(user_id)
        existing = store.collection.get(where={"content_hash": content_hash})
        if existing and existing.get("ids"):
            doc_id = existing["metadatas"][0]["doc_id"]
            return Document(
                id=doc_id,
                title=file_path.name,
                source_path=str(file_path),
                chunks=len(existing["ids"]),
                metadata={"dedup": True},
            )

        # Extract text
        text = self._extract_text(file_path)

        # Chunk (MD-aware if markdown)
        is_md = file_path.suffix.lower() == ".md"
        chunks = self._semantic_chunk(text, is_markdown=is_md)

        # Embed
        embeddings = self.encode(chunks)

        # Generate ID and save raw file
        doc_id = str(uuid.uuid4())
        user_dir = self.storage_root / user_id
        user_dir.mkdir(exist_ok=True)
        raw_path = user_dir / f"{doc_id}{file_path.suffix}"
        raw_path.write_bytes(content_bytes)

        # Store in vector DB
        store.add_document(
            doc_id=doc_id,
            chunks=chunks,
            embeddings=embeddings,
            metadata={
                "title": file_path.name,
                "source_path": str(file_path),
                "content_hash": content_hash,
            },
        )

        return Document(
            id=doc_id,
            title=file_path.name,
            source_path=str(raw_path),
            chunks=len(chunks),
            metadata={"original_path": str(file_path), "user_id": user_id},
        )

    def _extract_text(self, file_path: Path) -> str:
        suffix = file_path.suffix.lower()

        if suffix == ".pdf":
            import pdfplumber

            with pdfplumber.open(file_path) as pdf:
                return "\n\n".join(page.extract_text() or "" for page in pdf.pages)

        elif suffix in [".txt", ".md"]:
            return file_path.read_text(encoding="utf-8")

        elif suffix == ".docx":
            from docx import Document as DocxDocument

            doc = DocxDocument(file_path)
            return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())

        elif suffix == ".html":
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(file_path.read_text(), "html.parser")
            return soup.get_text(separator="\n", strip=True)

        else:
            raise ValueError(f"Unsupported format: {suffix}")

    def _semantic_chunk(
        self,
        text: str,
        target_size: int = 512,
        overlap: int = 64,
        is_markdown: bool = False,
    ) -> List[str]:
        """
        Semantic chunking by paragraphs.
        For Markdown: respects header boundaries (# ## ###) for better semantic chunks.
        """
        import re

        if is_markdown:
            # Split by headers for MD files
            # Pattern: lines starting with # (but not inside code blocks)
            sections = []
            current_section = []
            in_code_block = False

            for line in text.split("\n"):
                if line.strip().startswith("```"):
                    in_code_block = not in_code_block

                if not in_code_block and re.match(r"^#{1,6}\s", line):
                    if current_section:
                        sections.append("\n".join(current_section))
                        current_section = []

                current_section.append(line)

            if current_section:
                sections.append("\n".join(current_section))

            # Further split large sections
            chunks = []
            for section in sections:
                if len(section) > target_size:
                    # Split by paragraphs within section
                    paras = [
                        p.strip() for p in re.split(r"\n\s*\n", section) if p.strip()
                    ]
                    current = []
                    current_len = 0

                    for para in paras:
                        if current_len + len(para) > target_size and current:
                            chunks.append("\n\n".join(current))
                            current = [para]
                            current_len = len(para)
                        else:
                            current.append(para)
                            current_len += len(para)

                    if current:
                        chunks.append("\n\n".join(current))
                else:
                    chunks.append(section)

            return chunks if chunks else [text[:target_size]]

        else:
            # Regular paragraph chunking for non-MD
            paragraphs = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]

            chunks = []
            current = []
            current_len = 0

            for para in paragraphs:
                para_len = len(para)
                if current_len + para_len > target_size and current:
                    chunks.append("\n\n".join(current))
                    current = [para]
                    current_len = para_len
                else:
                    current.append(para)
                    current_len += para_len

            if current:
                chunks.append("\n\n".join(current))

            return chunks if chunks else [text[:target_size]]
